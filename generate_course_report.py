# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("个人学习笔记管理系统-课程大作业说明书.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 14 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, title in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, title, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(11)
        text = text[len(bold_prefix):]
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        styles[name].font.name = "宋体"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor(31, 77, 120)


def add_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("全栈开发课程大作业说明书")
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(26)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("题目：个人学习笔记管理系统")
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(16)

    for _ in range(5):
        doc.add_paragraph()
    for text in [
        "组号：______________________________",
        "姓名：季子皓（主要负责后端接口与业务逻辑）",
        "学号：______________________________",
        "班级：______________________________",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)
    doc.add_page_break()


def add_project_overview(doc):
    add_heading(doc, "一、项目概述", 1)
    add_heading(doc, "1.1 项目背景及目标", 2)
    add_para(doc, "在日常学习过程中，学生会不断产生课程笔记、代码片段、复习资料和问题记录。如果这些资料分散存放在纸质本、聊天文件或零散文档中，后期检索、归纳和复盘的效率会明显下降。个人学习笔记管理系统面向学生个人知识管理场景，采用 Web 前后端分离架构，将笔记编辑、分类整理、标签管理、检索筛选、数据导出和学习行为统计整合到统一平台中。")
    add_para(doc, "本项目的建设目标是：实现用户注册登录和身份认证，保证个人笔记数据隔离；实现笔记的新增、编辑、查看、收藏、删除、恢复和永久删除；支持分类、标签、关键词、日期和星标等多维检索；通过智能标签提取提升笔记归纳效率；提供 Markdown、TXT 和分类 ZIP 导出能力；通过统计看板展示分类占比、创作趋势、高频搜索词和活跃时间段，为用户复盘学习习惯提供依据。")

    add_heading(doc, "1.2 选题理由", 2)
    add_para(doc, "选择个人学习笔记管理系统作为全栈开发课程大作业，既贴近学生日常学习场景，又能够覆盖完整的前端、后端、数据库和部署流程。项目业务边界清晰，但功能点具有一定深度：后端需要完成认证授权、复杂条件查询、多表关系维护、文件下载、头像上传和统计聚合；前端需要完成富文本编辑、自动保存、筛选交互、图表展示和主题切换。因此该选题适合作为综合性全栈实践。")

    add_heading(doc, "1.3 分组说明", 2)
    add_para(doc, "小组成员包括方嘉辉、周子竣、季子皓。本人主要负责后端接口与业务逻辑实现，包括用户认证、笔记核心接口、分类与标签接口、统计聚合接口、导出接口、跨域与资源映射配置等，同时参与数据库字段设计、接口联调和报告撰写。")
    add_table(doc, ["成员", "主要分工", "具体内容"], [
        ["方嘉辉", "前端开发与页面交互", "负责 Vue3 页面搭建、Element Plus 组件使用、路由组织、Axios 请求封装、笔记列表与编辑页面交互。"],
        ["周子竣", "数据库设计、智能标签与测试", "负责 MySQL 表结构设计、数据字典整理、标签功能联调、系统功能测试和部署验证。"],
        ["季子皓", "后端接口与业务逻辑", "负责 Spring Boot REST API、JWT 认证、业务校验、JPA 查询、标签关系维护、导出下载和统计接口。"],
    ], widths=[3, 4, 9])

    add_heading(doc, "1.4 提交文档清单", 2)
    add_table(doc, ["序号", "提交内容", "说明"], [
        ["1", "前端项目源码", "frontend 目录，包含 Vue3、Vite、Element Plus、Axios、ECharts、WangEditor 相关代码。"],
        ["2", "后端项目源码", "backend 目录，包含 Spring Boot 控制器、实体、仓库、服务、拦截器和配置文件。"],
        ["3", "数据库 SQL 文件", "note_system.sql，用于初始化 users、notes、categories、tags、note_tags、search_logs 等表。"],
        ["4", "课程大作业说明书", "本文档，说明需求、设计、接口、实现、部署、总结和附件。"],
        ["5", "运行说明文档", "readme.md 和部署说明，描述环境依赖、启动方式和访问地址。"],
    ], widths=[1.5, 4.5, 10])


def add_requirements(doc):
    add_heading(doc, "二、需求分析", 1)
    add_heading(doc, "2.1 功能性需求", 2)
    add_para(doc, "系统采用普通用户单角色模式。用户登录后只能访问和管理自己的笔记、分类、标签、搜索日志和个人资料，所有核心接口均围绕当前认证用户展开。")
    add_table(doc, ["功能模块", "功能点", "需求说明"], [
        ["用户管理", "注册", "用户通过用户名、邮箱和密码创建账号，后端校验用户名唯一性并使用 BCrypt 加密保存密码。"],
        ["用户管理", "登录", "用户输入用户名和密码后，后端校验密码并签发 JWT，前端保存 Token 作为后续请求凭证。"],
        ["用户管理", "密码重置", "通过用户名和注册邮箱进行身份校验，校验成功后重置并重新加密密码。"],
        ["用户管理", "个人信息维护", "支持修改昵称、个性签名和上传头像，头像通过静态资源映射提供访问。"],
        ["笔记管理", "创建与编辑", "支持标题、富文本内容、纯文本内容、分类和标签保存，编辑页具备自动保存能力。"],
        ["笔记管理", "查看详情", "根据笔记 ID 查询详情，同时返回绑定标签，并更新最近访问时间。"],
        ["笔记管理", "回收站", "删除时先执行软删除，状态变为回收站；支持恢复和永久删除。"],
        ["星标收藏", "收藏切换", "支持对重要笔记进行星标标记，列表按星标和更新时间排序。"],
        ["分类管理", "分类维护", "支持创建、查询、修改、删除分类；删除分类后笔记分类字段置空，避免笔记丢失。"],
        ["标签管理", "标签维护", "保存笔记时同步维护 note_tags 关系，支持按用户获取标签列表和按标签筛选笔记。"],
        ["智能标签", "关键词提取", "后端根据笔记纯文本内容进行清洗、分词近似统计和高频词排序，返回推荐标签。"],
        ["检索筛选", "组合查询", "支持关键词、分类、星标、日期范围、标签和状态组合检索。"],
        ["数据导出", "单篇导出", "支持将单篇笔记导出为 Markdown 或 TXT。"],
        ["数据导出", "分类批量导出", "支持按分类打包导出 ZIP 文件，便于本地备份。"],
        ["统计看板", "学习数据统计", "统计总笔记数、星标数、分类数、分类占比、近 7 天趋势、高频搜索词和 24 小时活跃分布。"],
    ], widths=[3, 3.5, 9.5])

    add_heading(doc, "2.2 非功能性需求", 2)
    add_bullets(doc, [
        "性能需求：常用接口应在正常网络环境下快速响应，笔记列表查询通过数据库条件过滤减少前端处理压力。",
        "安全性需求：密码不得明文存储，受保护接口统一校验 JWT；后端接口需要根据 userId 验证数据归属，避免越权访问。",
        "可靠性需求：删除笔记采用回收站机制，导出功能支持本地备份；异常请求返回明确状态码和提示信息。",
        "易用性需求：前端提供富文本编辑、分类筛选、标签筛选、日期筛选、星标入口和夜间模式，降低使用成本。",
        "可维护性需求：后端按 Controller、Service、Repository、Entity、Config、Interceptor 分层组织，便于扩展和定位问题。",
        "兼容性需求：系统面向主流 PC 浏览器，前端使用 Vite 和 Vue3 构建，后端提供标准 RESTful API。",
    ])


def add_design(doc):
    add_heading(doc, "三、系统设计", 1)
    add_heading(doc, "3.1 总体架构设计", 2)
    add_para(doc, "系统采用前后端分离架构。前端为 Vue3 单页应用，负责页面渲染、表单校验、富文本编辑、图表展示和用户交互；后端为 Spring Boot REST API，负责认证授权、业务校验、数据持久化、文件处理和统计聚合；数据库使用 MySQL 保存用户、笔记、分类、标签和搜索日志等业务数据。")
    add_table(doc, ["层次", "组成", "职责"], [
        ["表现层", "Vue3、Vite、Element Plus、WangEditor、ECharts", "提供登录、首页、编辑器、分类管理、个人设置、统计看板等用户界面。"],
        ["通信层", "Axios 请求封装", "统一设置 API 基础路径、请求超时、Authorization 请求头和 401 响应处理。"],
        ["接口层", "Spring Boot Controller", "暴露 /api/auth、/api/notes、/api/categories、/api/statistics 等 REST 接口。"],
        ["认证层", "AuthInterceptor、JwtUtil", "统一拦截受保护接口，解析 JWT 并将 userId 注入 request 上下文。"],
        ["业务层", "UserService 与控制器内业务方法", "完成注册登录、密码重置、笔记保存、标签绑定、导出和统计等业务逻辑。"],
        ["持久层", "Spring Data JPA Repository", "通过 JPA 方法和 Native Query 访问 MySQL。"],
        ["数据层", "MySQL 8", "保存用户、笔记、分类、标签、笔记标签关联和搜索日志。"],
    ], widths=[2.5, 5, 8.5])
    add_para(doc, "后端请求链路为：浏览器页面发起请求，Axios 自动携带 Bearer Token；Spring Boot 的 AuthInterceptor 对 /api/** 进行拦截，登录和注册接口除外；Token 解析成功后将 userId 写入 request；Controller 获取当前 userId，调用 Repository 查询或更新当前用户数据；最终返回 JSON 或文件流给前端。")

    add_heading(doc, "3.2 技术选型", 2)
    add_table(doc, ["类别", "技术", "选择理由"], [
        ["后端语言与框架", "Java 17、Spring Boot 3.2.4", "生态成熟，适合快速构建 RESTful API，便于集成 Web、JPA、拦截器和配置管理。"],
        ["数据访问", "Spring Data JPA、Native Query", "普通 CRUD 通过 JpaRepository 简化开发，复杂组合筛选使用原生 SQL 保证灵活性。"],
        ["数据库", "MySQL 8", "关系模型稳定，适合保存用户、笔记、分类、标签等结构化数据。"],
        ["认证安全", "JJWT、BCrypt", "JWT 支持无状态认证，BCrypt 可对密码进行强哈希存储。"],
        ["前端框架", "Vue 3、Vite", "开发效率高，组件化能力强，适合单页应用。"],
        ["UI 与图表", "Element Plus、ECharts、WangEditor", "快速实现管理类界面、可视化看板和富文本编辑。"],
        ["部署配置", "环境变量、Dockerfile、Vercel/后端云服务配置", "数据库连接、上传目录和前端地址可随环境变化配置。"],
    ], widths=[3, 4.5, 8.5])

    add_heading(doc, "3.3 功能模块设计", 2)
    add_table(doc, ["模块", "后端入口", "核心职责"], [
        ["用户认证模块", "UserController、UserService、JwtUtil", "完成注册、登录、密码重置、Token 生成与解析、个人资料维护。"],
        ["笔记管理模块", "NoteController、NoteRepository", "完成笔记 CRUD、详情访问、最近访问、星标、软删除、恢复和永久删除。"],
        ["分类模块", "CategoryController、CategoryRepository", "完成分类列表、创建、修改和删除，并保证分类归属当前用户。"],
        ["标签模块", "TagRepository、NoteController", "完成用户标签查询、笔记标签绑定、标签筛选和智能推荐标签。"],
        ["检索模块", "NoteRepository.findByFilters", "支持关键词、分类、星标、日期、标签、回收站状态组合筛选。"],
        ["导出模块", "NoteController.exportSingleNote/exportCategoryNotes", "将笔记内容转换为文件流，支持单篇和分类 ZIP 导出。"],
        ["统计模块", "StatisticsController、SearchLogRepository", "聚合笔记数量、分类占比、趋势、高频搜索词和活跃时间段。"],
        ["系统配置模块", "WebConfig、application.yml", "配置跨域、拦截器、上传目录、静态资源映射和数据库连接。"],
    ], widths=[3, 5, 8])

    add_heading(doc, "3.4 数据库设计", 2)
    add_para(doc, "数据库围绕用户私有笔记管理进行设计，核心表包括 users、notes、categories、tags、note_tags 和 search_logs。用户与笔记、分类、标签均为一对多关系，笔记与标签为多对多关系，通过 note_tags 关联表实现。")
    add_table(doc, ["表名", "关键字段", "说明"], [
        ["users", "id、username、email、password、nickname、avatar、signature", "保存用户账号、密码哈希和个人资料。"],
        ["notes", "id、user_id、category_id、title、content、content_text、is_starred、status、last_accessed_at", "保存笔记主体数据，status 表示正常或回收站状态。"],
        ["categories", "id、user_id、name、parent_id", "保存用户自定义分类，parent_id 为后续层级分类扩展预留。"],
        ["tags", "id、user_id、name", "保存用户标签字典，避免不同用户标签混用。"],
        ["note_tags", "note_id、tag_id", "保存笔记和标签的多对多关系，联合主键避免重复绑定。"],
        ["search_logs", "id、user_id、keyword、created_at", "记录用户关键词搜索行为，用于统计高频搜索词。"],
    ], widths=[3, 6, 7])
    add_table(doc, ["关系", "基数", "设计说明"], [
        ["users -> notes", "1:N", "一个用户可拥有多篇笔记，删除用户时笔记级联删除。"],
        ["users -> categories", "1:N", "分类归属于用户，防止跨用户分类混用。"],
        ["users -> tags", "1:N", "标签归属于用户，同名标签在不同用户之间互不影响。"],
        ["categories -> notes", "1:N", "一类可包含多篇笔记；删除分类时笔记 category_id 置空。"],
        ["notes <-> tags", "N:M", "通过 note_tags 维护笔记和标签关系。"],
        ["users -> search_logs", "1:N", "搜索日志按用户隔离，用于个人统计分析。"],
    ], widths=[4, 2, 10])

    add_heading(doc, "3.5 接口与交互设计", 2)
    add_para(doc, "后端 API 统一以 /api 为前缀。除登录、注册外，接口都需要携带 Authorization: Bearer <token>。文件下载场景同时支持通过 URL token 参数传递 Token，便于 window.open 打开下载链接。")
    add_table(doc, ["方法", "接口", "主要参数", "功能说明"], [
        ["POST", "/api/auth/register", "username、password、email", "注册用户，校验用户名唯一并加密保存密码。"],
        ["POST", "/api/auth/login", "username、password", "登录校验，成功后返回 JWT 和用户基础信息。"],
        ["POST", "/api/auth/reset-password", "username、email、newPassword", "通过注册邮箱重置密码。"],
        ["PUT", "/api/auth/profile", "nickname、signature", "更新当前用户个人资料。"],
        ["POST", "/api/auth/avatar", "file", "上传头像文件并更新用户 avatar 地址。"],
        ["GET", "/api/notes", "status、keyword、categoryId、isStarred、startDate、endDate、tagName", "组合查询当前用户笔记列表。"],
        ["POST", "/api/notes", "title、content、contentText、categoryId、tags", "创建笔记并绑定标签。"],
        ["GET", "/api/notes/{id}", "id", "获取笔记详情和标签，并更新最近访问时间。"],
        ["PUT", "/api/notes/{id}", "title、content、contentText、categoryId、tags", "更新笔记主体和标签关系。"],
        ["DELETE", "/api/notes/{id}", "id", "软删除笔记，移入回收站。"],
        ["PUT", "/api/notes/{id}/restore", "id", "从回收站恢复笔记。"],
        ["DELETE", "/api/notes/{id}/permanent", "id", "永久删除当前用户笔记。"],
        ["PUT", "/api/notes/{id}/star", "id", "切换星标状态。"],
        ["GET", "/api/notes/recent", "无", "查询最近访问的 8 篇笔记。"],
        ["GET", "/api/notes/tags", "无", "查询当前用户所有标签。"],
        ["POST", "/api/notes/extract-tags", "text", "根据正文文本返回推荐标签。"],
        ["GET", "/api/notes/{id}/export", "type=md/txt", "导出单篇笔记。"],
        ["GET", "/api/notes/export/category", "categoryId、type", "按分类批量打包导出 ZIP。"],
        ["GET", "/api/categories", "无", "查询分类列表。"],
        ["POST", "/api/categories", "name、parentId", "新增分类。"],
        ["PUT", "/api/categories/{id}", "name", "修改分类名称。"],
        ["DELETE", "/api/categories/{id}", "id", "删除分类。"],
        ["GET", "/api/statistics", "无", "获取统计看板数据。"],
    ], widths=[2, 5, 4.5, 4.5])


def add_ui_and_implementation(doc):
    add_heading(doc, "四、界面与用户体验设计", 1)
    add_para(doc, "前端页面以学习笔记管理为核心工作流，主要页面包括登录注册页、首页笔记列表页、笔记编辑页、数据看板页和个人设置弹窗。由于本文档重点说明本人负责的后端接口与业务逻辑，页面说明主要用于解释接口交互来源。")
    add_table(doc, ["页面", "功能说明", "主要后端交互"], [
        ["登录/注册/重置页", "提供登录、注册、忘记密码三种模式，成功登录后保存 token 和用户信息。", "/auth/login、/auth/register、/auth/reset-password"],
        ["首页笔记列表", "展示全部、星标、分类、标签、回收站和最近访问笔记，支持关键词和日期筛选。", "/notes、/notes/recent、/notes/tags、/categories"],
        ["笔记编辑页", "提供标题、分类、标签和富文本编辑，支持 3 秒防抖自动保存和智能标签提取。", "/notes、/notes/{id}、/notes/extract-tags"],
        ["分类管理弹窗", "支持新增、重命名和删除分类。", "/categories"],
        ["个人设置弹窗", "支持头像上传、昵称和个性签名修改。", "/auth/avatar、/auth/profile"],
        ["数据看板", "展示总量卡片、分类占比、创作趋势、高频搜索词和活跃时段图表。", "/statistics"],
    ], widths=[3, 7, 6])

    add_heading(doc, "五、系统实现与部署", 1)
    add_heading(doc, "5.1 开发环境与环境配置", 2)
    add_table(doc, ["环境项", "版本或配置", "说明"], [
        ["JDK", "17+", "后端使用 Java 17 和 Spring Boot 3。"],
        ["后端框架", "Spring Boot 3.2.4", "提供 Web、JPA、配置、拦截器等能力。"],
        ["构建工具", "Maven", "backend/pom.xml 管理后端依赖。"],
        ["数据库", "MySQL 8.0+", "执行 note_system.sql 初始化表结构。"],
        ["前端运行时", "Node.js 18+", "运行 Vite 开发服务器和前端构建。"],
        ["前端框架", "Vue 3.4、Vite 5、Element Plus 2.6", "提供单页应用和组件库。"],
        ["默认端口", "后端 8080，前端 5173", "前端通过 /api 代理或环境变量访问后端。"],
    ], widths=[3.5, 4.5, 8])

    add_heading(doc, "5.2 项目结构与说明", 2)
    add_para(doc, "项目根目录采用前后端分离结构，主要目录如下：")
    add_table(doc, ["路径", "职责说明"], [
        ["backend/src/main/java/com/example/notesystem/controller", "REST 控制器，负责接收请求、获取当前用户、调用业务逻辑并返回响应。"],
        ["backend/src/main/java/com/example/notesystem/service", "业务服务层，目前 UserService 封装注册、登录、密码重置逻辑。"],
        ["backend/src/main/java/com/example/notesystem/repository", "JPA 仓库层，封装用户、笔记、分类、标签、统计查询。"],
        ["backend/src/main/java/com/example/notesystem/entity", "JPA 实体类，对应数据库表结构。"],
        ["backend/src/main/java/com/example/notesystem/interceptor", "认证拦截器，统一解析和校验 JWT。"],
        ["backend/src/main/java/com/example/notesystem/config", "Web 配置，包括跨域、静态资源映射和接口拦截范围。"],
        ["backend/src/main/resources/application.yml", "后端端口、数据库连接、JPA、上传目录和图片访问域名配置。"],
        ["frontend/src/views", "前端页面，包括 Login、Home、EditNote。"],
        ["frontend/src/components", "前端组件，包括 NoteEditor 和 DataStatistics。"],
        ["frontend/src/utils/request.js", "Axios 封装，统一添加 Token 和处理响应错误。"],
        ["note_system.sql", "数据库初始化脚本。"],
    ], widths=[7, 9])

    add_heading(doc, "5.3 本人负责的后端核心实现", 2)
    add_para(doc, "（1）用户认证与权限控制：UserService 使用 BCrypt 对密码进行哈希，登录成功后由 JwtUtil 生成包含 userId 和 username 的 JWT。AuthInterceptor 拦截 /api/** 请求，放行登录、注册和 OPTIONS 预检请求，解析 Token 后将 userId 注入 request，Controller 通过 userId 限制数据访问范围。")
    add_para(doc, "（2）笔记核心业务：NoteController 实现笔记创建、编辑、详情、列表、软删除、恢复、永久删除和星标切换。创建和编辑时同时保存富文本 content 与纯文本 contentText，既满足展示需要，又便于关键词检索和导出纯文本。")
    add_para(doc, "（3）组合检索：NoteRepository.findByFilters 使用原生 SQL 左连接 note_tags 和 tags，支持关键词、分类、星标、日期、标签、状态的组合筛选。关键词搜索时同时记录 search_logs，为统计看板提供数据来源。")
    add_para(doc, "（4）标签关系维护：保存笔记时先清空 note_tags 旧关系，再对传入标签逐个查找或创建 tags 记录，最后插入 note_tags 关系。该方式保证编辑笔记后标签关系与前端表单保持一致。")
    add_para(doc, "（5）导出与备份：单篇导出接口根据 type 返回 Markdown 或 TXT；分类导出接口查询某分类下的正常笔记，并使用 ZipOutputStream 在内存中生成 ZIP 压缩包。文件名使用 UTF-8 编码处理，兼容中文标题。")
    add_para(doc, "（6）统计聚合：StatisticsController 聚合当前用户正常笔记和分类数据，计算总笔记数、星标数、分类数、分类占比、近 7 天创作趋势，并结合 SearchLogRepository 查询高频搜索词和活跃时间分布。")
    add_para(doc, "（7）上传与跨域配置：WebConfig 将 /uploads/** 映射到配置的本地上传目录，头像上传后返回可访问 URL；同时配置 CORS，允许本地前端和部署环境前端访问后端接口。")

    add_heading(doc, "5.4 遇到的难点与解决方案", 2)
    add_table(doc, ["难点", "问题表现", "解决方案"], [
        ["JWT 与跨域联调", "前端跨域 POST 请求会先发送 OPTIONS 预检请求，如果认证拦截器拦截预检会导致 403。", "在 AuthInterceptor 中放行 OPTIONS 请求，并在 WebConfig 中统一配置 CORS、允许 Authorization 请求头。"],
        ["多条件检索与标签筛选", "笔记列表既要按标题正文模糊搜索，又要按分类、星标、日期和标签组合筛选，普通 JPA 方法名难以表达。", "在 NoteRepository 中使用 Native Query，LEFT JOIN 标签关系表，并通过参数为空时跳过条件的方式实现组合查询。"],
        ["笔记与标签多对多同步", "编辑笔记时标签可能新增、删除或重命名，如果只追加关系会产生脏数据。", "保存时先删除该笔记旧的 note_tags 关系，再查找或创建标签并重新绑定，保证关系表和当前提交一致。"],
        ["文件下载鉴权", "浏览器 window.open 下载文件时不方便设置 Authorization Header。", "认证拦截器除 Header 外也支持从 URL token 参数读取 Token，兼容下载场景。"],
        ["统计看板数据来源", "图表需要多种聚合数据，前端单独查询会增加请求次数和耦合度。", "后端 /api/statistics 一次性聚合返回看板所需数据，前端只负责图表渲染。"],
    ], widths=[3.5, 6, 6.5])

    add_heading(doc, "5.5 部署与运行说明", 2)
    add_bullets(doc, [
        "数据库准备：创建 note_system 数据库，执行根目录 note_system.sql 初始化表结构。",
        "后端配置：设置 DB_HOST、DB_PORT、DB_NAME、DB_USER、DB_PASSWORD、UPLOAD_PATH、IMAGE_BASE_URL、FRONTEND_URL 等环境变量。",
        "后端启动：进入 backend 目录，执行 mvn spring-boot:run，默认监听 http://localhost:8080。",
        "前端启动：进入 frontend 目录，执行 npm install 后运行 npm run dev，默认访问 http://localhost:5173。",
        "生产构建：前端执行 npm run build 生成 dist；后端可通过 Maven 打包或 Dockerfile 部署。",
        "测试账号：系统支持用户自行注册，首次使用可在登录页切换到注册模式创建账号。",
    ])

    add_heading(doc, "5.6 测试与验证说明", 2)
    add_para(doc, "项目包含 Spring Boot 默认上下文加载测试 DemoApplicationTests。功能验证主要通过前后端联调完成，覆盖用户注册登录、Token 过期处理、笔记新增编辑、组合筛选、分类管理、标签提取、回收站、导出和统计看板等流程。后续可进一步补充 Controller 层接口测试、Repository 查询测试和端到端自动化测试。")


def add_summary(doc):
    add_heading(doc, "六、总结与展望", 1)
    add_heading(doc, "6.1 项目完成情况", 2)
    add_table(doc, ["类别", "完成情况"], [
        ["已完成", "用户注册、登录、密码重置、JWT 认证、个人资料维护、头像上传。"],
        ["已完成", "笔记创建、编辑、查看、最近访问、星标、软删除、恢复和永久删除。"],
        ["已完成", "分类管理、标签管理、智能标签提取、组合检索和标签筛选。"],
        ["已完成", "单篇 Markdown/TXT 导出、分类 ZIP 批量导出。"],
        ["已完成", "统计看板数据接口，包括分类占比、近 7 天趋势、高频搜索词和活跃时段。"],
        ["待完善", "补充更完整的单元测试、接口自动化测试和异常场景测试。"],
        ["待完善", "进一步增强智能标签算法，例如接入分词库或模型服务，提高中文关键词质量。"],
        ["待完善", "完善线上部署脚本、日志监控、文件类型校验和接口参数校验。"],
    ], widths=[3, 13])

    add_heading(doc, "6.2 个人收获", 2)
    add_para(doc, "通过本次课程大作业，我对全栈项目从需求到落地的完整链路有了更系统的理解。作为主要负责后端接口与业务逻辑的成员，我重点实践了 Spring Boot REST API 设计、JWT 无状态认证、BCrypt 密码加密、JPA 数据访问、复杂 SQL 查询、多表关系维护、文件下载和统计聚合等内容。")
    add_para(doc, "在接口联调过程中，我也更深刻地体会到前后端约定的重要性，例如 Token 传递方式、响应数据结构、日期参数格式、下载接口鉴权方式和错误提示格式等。后端不仅要完成数据增删改查，还要考虑用户数据隔离、异常处理、接口易用性和后续扩展性。")
    add_para(doc, "后续如果继续完善该系统，我希望在三个方向上继续优化：第一，补充接口自动化测试和更严格的参数校验；第二，改进智能标签算法和全文检索能力；第三，完善部署、日志和监控，让系统更接近真实生产项目。")

    add_heading(doc, "七、附件清单", 1)
    add_table(doc, ["附件编号", "附件名称", "说明"], [
        ["附件1", "前端项目源码或打包文件", "frontend 目录或前端 dist 构建产物。"],
        ["附件2", "后端项目源码或打包文件", "backend 目录或 Maven 打包后的 jar 文件。"],
        ["附件3", "数据库导出 SQL 文件", "note_system.sql。"],
        ["附件4", "部署说明文档", "包含环境变量、数据库初始化、前后端启动和访问地址。"],
        ["附件5", "数据库数据字典", "可根据 note_system.sql 或 Navicat 数据字典导出。"],
    ], widths=[3, 5, 8])


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("个人学习笔记管理系统课程大作业说明书")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)


def main():
    doc = Document()
    set_doc_styles(doc)
    add_title_page(doc)
    add_project_overview(doc)
    add_requirements(doc)
    add_design(doc)
    add_ui_and_implementation(doc)
    add_summary(doc)
    add_footer(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
