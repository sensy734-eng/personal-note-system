# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_course_report import (
    add_bullets,
    add_footer,
    add_heading,
    add_para,
    add_table,
    set_doc_styles,
)


OUT = Path("个人学习笔记管理系统-课程大作业说明书-前端开发与页面交互.docx")


def add_frontend_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("全栈开发课程大作业说明书")
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(26)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("题目：个人学习笔记管理系统")
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(16)

    for _ in range(5):
        doc.add_paragraph()
    for text in [
        "组号：______________________________",
        "姓名：方嘉辉（主要负责前端开发与页面交互）",
        "学号：______________________________",
        "班级：______________________________",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)
    doc.add_page_break()


def add_overview(doc):
    add_heading(doc, "一、项目概述", 1)
    add_heading(doc, "1.1 项目背景及目标", 2)
    add_para(doc, "个人学习笔记管理系统面向学生学习资料整理场景，解决课程笔记分散、查找低效、复盘困难等问题。系统采用前后端分离架构，前端使用 Vue3、Vite、Element Plus、WangEditor、ECharts 和 Axios 构建单页应用，后端提供统一 REST API，数据库使用 MySQL 保存业务数据。")
    add_para(doc, "本项目的前端目标是：设计清晰的学习笔记工作台；实现登录、注册、忘记密码、笔记列表、笔记编辑、分类管理、标签筛选、个人设置、主题切换和统计看板等交互；通过 Axios 封装统一处理 Token 与异常提示；通过富文本编辑器和自动保存降低用户记录成本；通过图表组件展示学习数据。")

    add_heading(doc, "1.2 选题理由", 2)
    add_para(doc, "该选题贴近学生真实学习场景，既能体现管理类系统的基础交互，也能覆盖富文本编辑、筛选联动、图表可视化、文件下载、登录态维护和前后端接口联调等典型全栈开发内容。前端部分不仅是页面展示，还需要对复杂状态、异步请求、用户反馈和操作流程进行组织，因此具有较好的课程实践价值。")

    add_heading(doc, "1.3 分组说明", 2)
    add_para(doc, "小组成员包括方嘉辉、周子竣、季子皓。本人主要负责前端开发与页面交互，包括 Vue3 页面搭建、路由守卫、Element Plus 组件应用、Axios 请求封装、富文本编辑器集成、笔记筛选联动、分类与个人设置弹窗、数据看板展示、夜间模式和前后端接口联调。")
    add_table(doc, ["成员", "主要分工", "具体内容"], [
        ["方嘉辉", "前端开发与页面交互", "负责 Login、Home、EditNote、NoteEditor、DataStatistics 等页面和组件，完成路由、请求、状态、表单、图表和主题交互。"],
        ["周子竣", "数据库设计、智能标签与测试", "负责 MySQL 表结构、数据字典、标签功能联调、系统测试和部署验证。"],
        ["季子皓", "后端接口与业务逻辑", "负责 Spring Boot REST API、JWT 认证、业务校验、JPA 查询、导出下载和统计接口。"],
    ], widths=[3, 4, 9])

    add_heading(doc, "1.4 提交文档清单", 2)
    add_table(doc, ["序号", "提交内容", "说明"], [
        ["1", "前端项目源码", "frontend 目录，包含 Vue3 页面、组件、请求封装、路由入口和 Vite 配置。"],
        ["2", "后端项目源码", "backend 目录，提供前端页面调用的 REST API。"],
        ["3", "数据库 SQL 文件", "note_system.sql，用于初始化用户、笔记、分类、标签等数据表。"],
        ["4", "课程大作业说明书", "本文档，重点说明本人负责的前端开发与页面交互。"],
        ["5", "运行说明文档", "readme.md，说明依赖安装、前后端启动和访问方式。"],
    ], widths=[1.5, 4.5, 10])


def add_requirements(doc):
    add_heading(doc, "二、需求分析", 1)
    add_heading(doc, "2.1 功能性需求", 2)
    add_table(doc, ["功能模块", "前端页面或组件", "需求说明"], [
        ["用户管理", "Login.vue", "提供登录、注册、密码重置三种模式，完成表单校验、模式切换、成功提示和登录后跳转。"],
        ["登录态维护", "main.js、request.js", "通过路由守卫保护首页和编辑页；Axios 请求前自动携带 Bearer Token；401 时清理登录态并跳转登录页。"],
        ["首页工作台", "Home.vue", "以侧边栏、顶部筛选区和笔记卡片组成主界面，展示全部、星标、分类、标签、回收站和最近访问。"],
        ["笔记筛选", "Home.vue", "支持关键词、日期范围、分类、标签、星标和回收站状态联动筛选，并将条件转换为接口参数。"],
        ["笔记操作", "Home.vue", "支持新建、编辑、星标切换、软删除、恢复、永久删除、单篇导出和分类批量导出。"],
        ["分类管理", "Home.vue 弹窗", "支持新增分类、行内编辑分类名称、删除分类，并给出必要确认提示。"],
        ["个人设置", "Home.vue 弹窗", "支持头像上传、昵称和个性签名修改，并同步 localStorage 中的用户信息。"],
        ["富文本编辑", "NoteEditor.vue", "集成 WangEditor v5，提供工具栏、HTML 内容双向绑定和纯文本内容回传。"],
        ["编辑页面", "EditNote.vue", "支持标题、分类、标签、富文本内容编辑，区分新建和编辑状态。"],
        ["自动保存", "EditNote.vue", "监听表单变化，使用 3 秒防抖自动保存，减少用户忘记保存导致的数据丢失。"],
        ["智能标签", "EditNote.vue", "调用标签提取接口，将返回结果去重后合并到当前标签选择中。"],
        ["数据看板", "DataStatistics.vue", "调用统计接口，使用 ECharts 展示总量卡片、饼图、折线图、柱状图和时段分布图。"],
        ["主题切换", "App.vue、Home.vue", "支持亮色/暗色模式切换，并将偏好保存到 localStorage。"],
    ], widths=[3, 4, 9])

    add_heading(doc, "2.2 非功能性需求", 2)
    add_bullets(doc, [
        "易用性：页面布局应符合笔记管理习惯，常用操作入口明显，编辑和筛选路径清晰。",
        "响应性：接口请求期间要显示加载状态，操作成功或失败要通过 Element Plus 消息组件及时反馈。",
        "一致性：按钮、弹窗、表格、卡片、菜单和图表均使用 Element Plus 与 ECharts 统一风格。",
        "安全性：前端不得访问受保护页面时绕过登录态，所有业务请求应统一携带 Token。",
        "可维护性：页面、组件和工具函数分离，富文本编辑器与统计看板封装为可复用组件。",
        "兼容性：面向主流 PC 浏览器运行，使用 Vite 开发服务器和标准 Vue Router 单页路由。",
    ])


def add_design(doc):
    add_heading(doc, "三、系统设计", 1)
    add_heading(doc, "3.1 前端总体架构设计", 2)
    add_para(doc, "前端采用 Vue3 单页应用架构，入口文件 main.js 创建应用、注册 Vue Router 和 Element Plus。App.vue 作为顶层容器，通过 router-view 渲染页面，并在挂载时读取本地主题偏好。页面层包括 Login.vue、Home.vue 和 EditNote.vue；组件层包括 NoteEditor.vue 和 DataStatistics.vue；工具层通过 request.js 封装 Axios。")
    add_table(doc, ["层次", "文件或技术", "职责"], [
        ["应用入口", "main.js", "创建 Vue 应用，配置路由、注册 Element Plus，定义登录态路由守卫。"],
        ["顶层容器", "App.vue", "承载 router-view，初始化暗色主题偏好。"],
        ["页面层", "Login.vue、Home.vue、EditNote.vue", "完成登录注册、首页工作台和笔记编辑三大业务页面。"],
        ["组件层", "NoteEditor.vue、DataStatistics.vue", "封装富文本编辑器和数据看板，降低页面复杂度。"],
        ["请求层", "utils/request.js", "统一配置 baseURL、超时、Token 注入、响应剥离和错误处理。"],
        ["UI 组件库", "Element Plus、icons-vue", "提供表单、菜单、弹窗、卡片、表格、上传、日期选择和图标。"],
        ["可视化与编辑器", "ECharts、WangEditor", "提供统计图表和富文本编辑能力。"],
    ], widths=[3, 5, 8])

    add_heading(doc, "3.2 页面与交互流程设计", 2)
    add_table(doc, ["页面", "核心布局", "交互流程"], [
        ["登录页", "居中卡片、动态表单、模式切换链接", "用户可在登录、注册、重置密码之间切换；提交前进行必填、长度和邮箱格式校验；登录成功后保存 token 并跳转首页。"],
        ["首页", "左侧菜单、顶部搜索与用户区、主内容笔记卡片", "用户通过菜单切换全部、星标、分类、标签、回收站和数据看板；筛选条件变化后刷新笔记列表。"],
        ["编辑页", "顶部操作栏、标题输入、分类标签栏、富文本编辑器", "新建时保存后自动替换到编辑路由；编辑时加载详情和标签；内容变化触发防抖自动保存。"],
        ["分类弹窗", "输入框、添加按钮、分类表格", "新增分类后刷新列表；点击编辑图标进入行内编辑；删除前弹出确认。"],
        ["个人设置弹窗", "头像上传、昵称输入、签名文本域", "上传头像时携带 Authorization 头；保存资料后更新本地用户信息和页面展示。"],
        ["统计看板", "三张统计卡片、四个图表区域", "页面挂载后请求统计数据，nextTick 后初始化图表，组件卸载时销毁图表实例。"],
    ], widths=[3, 5, 8])

    add_heading(doc, "3.3 前端接口交互设计", 2)
    add_para(doc, "前端请求统一通过 request.js 发起。开发环境下 baseURL 默认为 /api，并由 Vite 代理转发到 http://localhost:8080；生产环境可通过 VITE_API_BASE_URL 指定后端地址。请求拦截器从 localStorage 读取 token 并写入 Authorization；响应拦截器直接返回 response.data，减少页面层重复取值；当响应状态码为 401 时清理 token 和 userInfo，并跳转登录页。")
    add_table(doc, ["前端功能", "调用接口", "页面处理"], [
        ["登录", "POST /auth/login", "保存 token、userInfo，并跳转首页。"],
        ["注册", "POST /auth/register", "成功后切换回登录模式。"],
        ["重置密码", "POST /auth/reset-password", "成功后清空新密码并返回登录模式。"],
        ["笔记列表", "GET /notes", "根据当前菜单、搜索词和日期组装 params，更新 notesList。"],
        ["最近访问", "GET /notes/recent", "首页全部笔记视图展示最近访问卡片。"],
        ["标签列表", "GET /notes/tags", "侧边栏标签菜单展示。"],
        ["分类列表", "GET /categories", "侧边栏分类菜单、编辑页分类选择和分类管理表格复用。"],
        ["保存笔记", "POST/PUT /notes", "新建或更新笔记，手动保存后返回首页，自动保存只更新保存时间。"],
        ["智能标签", "POST /notes/extract-tags", "将推荐标签合并进 tags 数组并去重。"],
        ["统计看板", "GET /statistics", "渲染 ECharts 饼图、折线图、柱状图和时段分布图。"],
        ["头像上传", "POST /auth/avatar", "通过 el-upload 设置 headers，成功后更新头像 URL。"],
    ], widths=[3, 5, 8])

    add_heading(doc, "3.4 UI 与用户体验设计", 2)
    add_bullets(doc, [
        "首页采用管理台式布局，左侧导航固定功能入口，顶部集中放置搜索、日期筛选、主题切换和用户操作。",
        "笔记卡片展示标题、正文摘要、更新时间和常用操作，方便用户快速浏览与处理。",
        "回收站模式下卡片操作切换为恢复和永久删除，避免用户误把删除状态下的笔记当成普通笔记处理。",
        "编辑页把标题、分类、标签和正文放在同一编辑上下文中，减少跨页面操作。",
        "自动保存通过 lastSavedTime 给出反馈，降低用户对保存状态的不确定感。",
        "数据看板把数字卡片和图表组合起来，既能快速查看总量，也能观察分类结构和使用习惯。",
        "夜间模式通过 html.dark 与 Element Plus 暗色变量协同，适合长时间阅读和编辑。",
    ])


def add_implementation(doc):
    add_heading(doc, "四、系统实现与部署", 1)
    add_heading(doc, "4.1 开发环境与依赖", 2)
    add_table(doc, ["类别", "技术或版本", "说明"], [
        ["前端框架", "Vue 3.4.21", "采用 Composition API 组织页面状态和生命周期。"],
        ["构建工具", "Vite 5.2.0", "提供快速开发服务器、热更新和生产构建。"],
        ["组件库", "Element Plus 2.6.1", "用于表单、菜单、弹窗、表格、卡片、上传、消息提示等 UI。"],
        ["图标库", "@element-plus/icons-vue", "用于菜单、按钮和操作入口图标。"],
        ["HTTP 请求", "Axios 1.6.8", "封装统一请求实例，处理 Token 和错误响应。"],
        ["富文本编辑", "WangEditor v5", "用于笔记正文编辑，并获取 HTML 与纯文本内容。"],
        ["数据可视化", "ECharts 5.5.0", "用于统计看板图表展示。"],
        ["路由", "Vue Router 4.3.0", "定义登录页、首页、新建和编辑页路由，并配置路由守卫。"],
    ], widths=[3.5, 4.5, 8])

    add_heading(doc, "4.2 项目结构与说明", 2)
    add_table(doc, ["路径", "职责说明"], [
        ["frontend/src/main.js", "前端应用入口，注册路由和 Element Plus，定义 requireAuth 路由守卫。"],
        ["frontend/src/App.vue", "顶层组件，负责路由出口和主题偏好初始化。"],
        ["frontend/src/views/Login.vue", "登录、注册、重置密码表单页面。"],
        ["frontend/src/views/Home.vue", "首页工作台，包含菜单、筛选、笔记列表、分类管理、个人设置和主题切换。"],
        ["frontend/src/views/EditNote.vue", "笔记新建与编辑页面，包含自动保存、导出和智能标签交互。"],
        ["frontend/src/components/NoteEditor.vue", "WangEditor 封装组件，向父组件同步 HTML 和纯文本内容。"],
        ["frontend/src/components/DataStatistics.vue", "ECharts 图表组件，展示统计看板。"],
        ["frontend/src/utils/request.js", "Axios 实例封装，统一处理 Token、响应数据和错误提示。"],
        ["frontend/vite.config.js", "Vite 配置，设置开发端口 5173 和 /api 代理。"],
        ["frontend/package.json", "前端依赖和 npm scripts 管理。"],
    ], widths=[6, 10])

    add_heading(doc, "4.3 本人负责的前端核心实现", 2)
    add_para(doc, "（1）路由与登录态控制：在 main.js 中配置 /login、/、/note/new、/note/edit/:id 等路由，对首页和编辑页使用 beforeEnter 路由守卫，未登录用户自动跳转到登录页。")
    add_para(doc, "（2）请求封装与异常处理：在 request.js 中创建 Axios 实例，统一设置 baseURL 和 timeout。请求前自动读取 localStorage 中的 token 并写入 Authorization 请求头；响应成功时直接返回核心数据；当后端返回 401 时清理本地登录信息并跳转登录页。")
    add_para(doc, "（3）登录注册页面：Login.vue 使用同一套表单状态支持登录、注册和密码重置三种模式，借助 Element Plus 表单规则完成用户名、密码、邮箱和新密码校验，并通过地址 hash 支持刷新后保持注册或重置模式。")
    add_para(doc, "（4）首页工作台：Home.vue 负责笔记列表主界面，左侧菜单展示全部、星标、分类、标签、回收站和数据看板；顶部提供全文检索、日期范围筛选、分类导出、主题切换、用户下拉菜单和新建笔记入口。")
    add_para(doc, "（5）复杂筛选联动：fetchNotes 根据 currentFilter、searchQuery 和 dateRange 动态组装后端参数，例如 status、keyword、categoryId、tagName、isStarred、startDate 和 endDate，实现多个筛选条件与后端接口的统一联动。")
    add_para(doc, "（6）分类与个人设置交互：分类管理使用 el-dialog 和 el-table 实现新增、行内编辑和删除确认；个人设置弹窗集成头像上传、昵称和签名修改，上传时通过 headers 携带 Token，成功后同步 localStorage 和页面头像。")
    add_para(doc, "（7）富文本编辑与自动保存：EditNote.vue 使用 NoteEditor 组件接收 HTML 内容，同时通过 textChange 获取纯文本摘要；watch 深度监听 noteForm，使用 3 秒定时器实现防抖自动保存；手动保存时显示 loading 和成功提示。")
    add_para(doc, "（8）统计看板实现：DataStatistics.vue 在挂载时请求 /statistics，等待 DOM 更新后初始化 ECharts 实例，分别渲染分类占比饼图、创作趋势折线图、高频搜索词柱状图和 24 小时活跃分布图，组件卸载时销毁图表实例。")
    add_para(doc, "（9）夜间模式：Home.vue 中通过 el-switch 切换 html.dark 类名，并写入 localStorage；App.vue 首次挂载时读取 theme-mode，保证刷新后主题偏好仍然生效。")

    add_heading(doc, "4.4 遇到的难点与解决方案", 2)
    add_table(doc, ["难点", "问题表现", "解决方案"], [
        ["登录态统一维护", "多个页面都需要携带 Token，如果每个请求手动写 Header 容易遗漏。", "封装 Axios 请求拦截器，统一读取 localStorage token 并写入 Authorization。"],
        ["401 失效处理", "Token 过期后页面可能继续停留在业务页，导致后续操作连续报错。", "响应拦截器集中处理 401，清理 token 和 userInfo，并延迟跳转到登录页。"],
        ["筛选条件组合复杂", "首页需要同时支持菜单筛选、关键词搜索、日期范围、星标、标签和回收站状态。", "使用 currentFilter 作为核心状态，在 fetchNotes 中统一转换为后端接口参数。"],
        ["自动保存触发频繁", "富文本编辑过程中内容变化很频繁，如果每次都请求会造成接口压力。", "使用 watch + setTimeout 实现 3 秒防抖保存，并在组件卸载时清理计时器。"],
        ["新建笔记路由切换", "新建笔记第一次保存后需要拿到后端返回 ID，否则后续自动保存无法变成编辑模式。", "保存成功后读取 noteId，使用 router.replace 切换到 /note/edit/:id。"],
        ["图表初始化时机", "ECharts 需要 DOM 容器存在后才能初始化，否则图表尺寸异常。", "统计数据请求完成后使用 nextTick，再初始化图表实例。"],
        ["文件下载鉴权", "window.open 无法像 Axios 一样自动设置 Authorization Header。", "导出链接拼接 token 参数，配合后端兼容下载场景。"],
    ], widths=[3.5, 6, 6.5])

    add_heading(doc, "4.5 部署与运行说明", 2)
    add_bullets(doc, [
        "安装依赖：进入 frontend 目录，执行 npm install。",
        "开发运行：执行 npm run dev，默认访问 http://localhost:5173。",
        "接口代理：开发环境下 vite.config.js 将 /api 代理到 http://localhost:8080。",
        "生产构建：执行 npm run build 生成 dist 目录，可部署到 Vercel、Nginx 或其他静态托管服务。",
        "生产接口地址：通过 VITE_API_BASE_URL 配置后端 API 地址，避免写死本地地址。",
        "测试账号：系统支持在登录页注册新用户，注册完成后回到登录模式。"],
    )

    add_heading(doc, "4.6 测试与验证说明", 2)
    add_para(doc, "前端功能主要通过页面联调验证，包括登录注册、路由守卫、Token 自动携带、401 跳转、笔记列表筛选、新建编辑、自动保存、分类管理、标签筛选、头像上传、导出、统计看板和夜间模式。后续可以进一步补充组件单元测试和端到端自动化测试，例如使用 Vitest 测试工具函数，使用 Playwright 覆盖登录、编辑和筛选主流程。")


def add_summary(doc):
    add_heading(doc, "五、总结与展望", 1)
    add_heading(doc, "5.1 项目完成情况", 2)
    add_table(doc, ["类别", "完成情况"], [
        ["已完成", "登录、注册、密码重置页面及表单校验。"],
        ["已完成", "Vue Router 路由配置和登录态路由守卫。"],
        ["已完成", "Axios 请求封装、Token 注入、401 统一处理和错误提示。"],
        ["已完成", "首页笔记列表、最近访问、分类/标签/星标/回收站筛选。"],
        ["已完成", "分类管理弹窗、个人设置弹窗和头像上传交互。"],
        ["已完成", "富文本编辑器集成、纯文本回传、自动保存和智能标签提取。"],
        ["已完成", "统计看板图表展示和夜间模式持久化。"],
        ["待完善", "进一步优化移动端响应式布局和小屏适配。"],
        ["待完善", "补充前端组件单元测试和端到端自动化测试。"],
        ["待完善", "优化中文编码显示、下载地址配置和部分硬编码 localhost 地址。"],
    ], widths=[3, 13])

    add_heading(doc, "5.2 个人收获", 2)
    add_para(doc, "通过本次课程大作业，我对 Vue3 前端工程化开发和前后端分离项目的协作方式有了更完整的认识。作为主要负责前端开发与页面交互的成员，我实践了 Vue Router、Element Plus、Axios、WangEditor、ECharts、Vite 代理、localStorage 登录态维护和组件化开发等内容。")
    add_para(doc, "在联调过程中，我认识到前端不仅要把页面做出来，还需要把用户操作流程设计顺畅。例如登录态过期时要及时提示并跳转，保存笔记时要区分手动保存和自动保存，复杂筛选要统一映射为接口参数，上传头像和文件导出也要考虑鉴权方式。")
    add_para(doc, "如果后续继续完善系统，我希望进一步优化页面响应式布局、补充自动化测试、减少硬编码接口地址，并提升富文本编辑体验，让系统在真实学习场景中更加稳定易用。")

    add_heading(doc, "六、附件清单", 1)
    add_table(doc, ["附件编号", "附件名称", "说明"], [
        ["附件1", "前端项目源码或打包文件", "frontend 目录或 npm run build 生成的 dist。"],
        ["附件2", "后端项目源码或打包文件", "backend 目录或 Maven 打包后的 jar。"],
        ["附件3", "数据库导出 SQL 文件", "note_system.sql。"],
        ["附件4", "部署说明文档", "包含前端依赖安装、Vite 代理、生产构建和接口地址配置。"],
        ["附件5", "数据库数据字典", "可根据 note_system.sql 或 Navicat 导出。"],
    ], widths=[3, 5, 8])


def main():
    doc = Document()
    set_doc_styles(doc)
    add_frontend_title_page(doc)
    add_overview(doc)
    add_requirements(doc)
    add_design(doc)
    add_implementation(doc)
    add_summary(doc)
    add_footer(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
